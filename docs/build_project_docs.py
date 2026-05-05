from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCES = ROOT / "sources"
OUTPUTS = ROOT
ACCENT = RGBColor(30, 82, 152)
ACCENT_LIGHT = "DCE6F7"
TEXT_DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(92, 110, 133)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_DARK

    for style_name, size in [("Heading 1", 24), ("Heading 2", 16), ("Heading 3", 12.5)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT


def add_cover(document: Document, title: str, subtitle: str) -> None:
    for _ in range(4):
        document.add_paragraph("")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(8)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(10)
    run2 = p2.add_run(subtitle)
    run2.font.name = "Calibri"
    run2.font.size = Pt(13)
    run2.font.color.rgb = MUTED

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Proyecto SmartPark | Fecha {datetime.now().strftime('%d/%m/%Y')}")
    run3.font.name = "Calibri"
    run3.font.size = Pt(10.5)
    run3.font.color.rgb = MUTED

    document.add_page_break()


def add_footer(document: Document) -> None:
    for section in document.sections:
        section.footer.is_linked_to_previous = False
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("SmartPark - Documentacion del proyecto")
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def parse_markdown_tables(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start_index
    while i < len(lines):
        line = lines[i].rstrip()
        if "|" not in line:
            break
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    data_rows = [row for row in rows[2:] if any(cell.strip() for cell in row)]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        hdr_cells[idx].text = value
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[idx], "C9D7F0")
        for para in hdr_cells[idx].paragraphs:
            for run in para.runs:
                run.font.bold = True

    for row in data_rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx < len(cells):
                cells[idx].text = value
                cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    document.add_paragraph("")


def add_bullet(document: Document, text: str, *, level: int = 0, numbered: bool = False) -> None:
    p = document.add_paragraph()
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    prefix = "1. " if numbered else "• "
    run = p.add_run(prefix + text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def build_docx_from_markdown(source_path: Path, output_path: Path) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), source_path.stem.replace("_", " "))
    subtitle = next((line[3:].strip() for line in lines if line.startswith("## ")), "SmartPark")

    document = Document()
    set_page_margins(document)
    apply_base_styles(document)
    add_cover(document, title, subtitle)
    add_footer(document)

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw.strip()

        if not line:
            document.add_paragraph("")
            i += 1
            continue

        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            document.add_heading(clean_inline(line[3:].strip()), level=1)
            i += 1
            continue
        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:].strip()), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            document.add_heading(clean_inline(line[5:].strip()), level=3)
            i += 1
            continue

        if "|" in line and not line.startswith("```"):
            rows, next_index = parse_markdown_tables(lines, i)
            add_markdown_table(document, rows)
            i = next_index
            continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            for code_line in code_lines:
                run = p.add_run(code_line + "\n")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue

        bullet_match = re.match(r"^-\s+(.*)$", line)
        if bullet_match:
            add_bullet(document, clean_inline(bullet_match.group(1)))
            i += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            add_bullet(document, clean_inline(number_match.group(1)), numbered=True)
            i += 1
            continue

        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean_inline(line))
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    mapping = {
        SOURCES / "Manual_Tecnico_SmartPark.md": OUTPUTS / "Manual_Tecnico_SmartPark.docx",
        SOURCES / "Manual_Usuario_SmartPark.md": OUTPUTS / "Manual_Usuario_SmartPark.docx",
        SOURCES / "Acta_Proyecto_SmartPark.md": OUTPUTS / "Acta_Proyecto_SmartPark.docx",
    }
    for source, output in mapping.items():
        build_docx_from_markdown(source, output)
        print(output)


if __name__ == "__main__":
    main()
    def clean_inline(text: str) -> str:
        return text.replace("`", "")
